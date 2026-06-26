"""Generate Word doc with corrected FR/EN infographic translations."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROWS = [
    (1, "Carte Europe", "Titre principal", "Le marché des paiements par carte en Europe : volumes et influence internationale", "Europe's card payments market: volumes and international influence"),
    (2, "Carte Europe", "Logo", "IRIS", "IRIS"),
    (3, "Carte Europe", "Légende carte", "États de la zone euro", "Eurozone countries"),
    (4, "Carte Europe", "Légende volumes — titre", "Répartition de la valeur totale des transactions en 2024 par États, en milliards d'euros", "Distribution of total transaction value in 2024 by country, in EUR billions"),
    (5, "Carte Europe", "Légende volumes — échelle", "856,7 (R-U)", "856.7 (UK)"),
    (6, "Carte Europe", "Légende volumes — échelle", "500", "500"),
    (7, "Carte Europe", "Légende volumes — échelle", "inférieur à 100", "less than 100"),
    (8, "Carte Europe", "Légende volumes — données", "Absence de données", "No data available"),
    (9, "Carte Europe", "Légende dépendance — titre", "Niveau de dépendance à des réseaux internationaux, établi sur la part des réseaux internationaux empruntés par les paiements par carte", "Level of dependence on international networks, based on their share of card payments"),
    (10, "Carte Europe", "Légende dépendance — faible", "Faible (réseaux internationaux inférieur à 40 %)", "Low (international networks less than 40%)"),
    (11, "Carte Europe", "Légende dépendance — modéré", "Modéré (réseaux internationaux compris entre 40 et 75 %)", "Moderate (international networks between 40% and 75%)"),
    (12, "Carte Europe", "Légende dépendance — total", "Total (100 %)", "Total (100%)"),
    (13, "Carte Europe", "Pays — carte", "ROYAUME-UNI", "UNITED KINGDOM"),
    (14, "Carte Europe", "Pays — carte", "IRLANDE", "IRELAND"),
    (15, "Carte Europe", "Pays — carte", "NORVÈGE", "NORWAY"),
    (16, "Carte Europe", "Pays — carte", "SUÈDE", "SWEDEN"),
    (17, "Carte Europe", "Pays — carte", "FINLANDE", "FINLAND"),
    (18, "Carte Europe", "Pays — carte", "ESTONIE", "ESTONIA"),
    (19, "Carte Europe", "Pays — carte", "LETTONIE", "LATVIA"),
    (20, "Carte Europe", "Pays — carte", "LITUANIE", "LITHUANIA"),
    (21, "Carte Europe", "Pays — carte", "DANEMARK", "DENMARK"),
    (22, "Carte Europe", "Pays — carte", "PAYS-BAS", "NETHERLANDS"),
    (23, "Carte Europe", "Pays — carte", "BELGIQUE", "BELGIUM"),
    (24, "Carte Europe", "Pays — carte", "LUX.", "LUX."),
    (25, "Carte Europe", "Pays — carte", "FRANCE", "FRANCE"),
    (26, "Carte Europe", "Pays — carte", "ALLEMAGNE", "GERMANY"),
    (27, "Carte Europe", "Pays — carte", "AUTRICHE", "AUSTRIA"),
    (28, "Carte Europe", "Pays — carte", "SLOVAQUIE", "SLOVAKIA"),
    (29, "Carte Europe", "Pays — carte", "SLOVÉNIE", "SLOVENIA"),
    (30, "Carte Europe", "Pays — carte", "CROATIE", "CROATIA"),
    (31, "Carte Europe", "Pays — carte", "ITALIE", "ITALY"),
    (32, "Carte Europe", "Pays — carte", "MALTE", "MALTA"),
    (33, "Carte Europe", "Pays — carte", "CHYPRE", "CYPRUS"),
    (34, "Carte Europe", "Pays — carte", "BULGARIE", "BULGARIA"),
    (35, "Carte Europe", "Pays — carte", "GRÈCE", "GREECE"),
    (36, "Carte Europe", "Pays — carte", "ESPAGNE", "SPAIN"),
    (37, "Carte Europe", "Pays — carte", "PORTUGAL", "PORTUGAL"),
    (38, "Carte Europe", "Source", "Sources : données compilées à partir des rapports des banques centrales de chaque pays de la zone euro, des rapports annuels des schémas domestiques et internationaux", "Sources: data compiled from the reports of the central banks of each eurozone country, and from the annual reports of domestic and international schemes"),
    (39, "Alternatives souveraines", "Titre principal", "Alternatives souveraines aux dépendances étrangères au sein des chaînes de paiement européennes", "Sovereign alternatives to foreign dependencies within European payment chains"),
    (40, "Alternatives souveraines", "Logo", "IRIS", "IRIS"),
    (41, "Alternatives souveraines", "Section 1 — titre", "LE CLOUD", "THE CLOUD"),
    (42, "Alternatives souveraines", "Section 1 — acteurs US", "Amazon Web Services (AWS), Microsoft Azure, Google Cloud", "Amazon Web Services (AWS), Microsoft Azure, Google Cloud"),
    (43, "Alternatives souveraines", "Section 1 — stat US", "65 à 70 % du cloud européen concentré chez trois acteurs américains", "65–70% of the European cloud market is concentrated among three US providers."),
    (44, "Alternatives souveraines", "Section 1 — description US", "Hébergent les systèmes des banques, des acquéreurs et des réseaux de cartes européens.", "They host the systems of European banks, acquirers and card networks."),
    (45, "Alternatives souveraines", "Section 1 — acteurs EU", "Outscale, OVHcloud, Cloud Temple, S3NS, CEGEDIM, Orange", "Outscale, OVHcloud, Cloud Temple, S3NS, CEGEDIM, Orange"),
    (46, "Alternatives souveraines", "Section 1 — stat EU", "6 fournisseurs européens qualifiés SecNumCloud", "6 European providers are SecNumCloud-certified"),
    (47, "Alternatives souveraines", "Section 1 — description EU", "Le label SecNumCloud impose l'immunité aux lois extraterritoriales états-uniennes (CLOUD Act, FISA)", "The SecNumCloud label requires protection against US extraterritorial laws (CLOUD Act, FISA)"),
    (48, "Alternatives souveraines", "Section 2 — titre", "PAIEMENT", "PAYMENT"),
    (49, "Alternatives souveraines", "Section 2 — portefeuilles US", "Apple Pay, Google Pay", "Apple Pay, Google Pay"),
    (50, "Alternatives souveraines", "Section 2 — stat portefeuilles US", "15 % des paiements par carte de proximité en France réalisés par portefeuille mobile en 2024", "15% of in-person card payments in France were made via mobile wallets in 2024."),
    (51, "Alternatives souveraines", "Section 2 — description portefeuilles US", "Contrôlent l'interface de paiement sur smartphone", "They control the payment interface on smartphones"),
    (52, "Alternatives souveraines", "Section 2 — alternative EU", "WERO", "WERO"),
    (53, "Alternatives souveraines", "Section 2 — stat WERO", "53 millions d'utilisateurs en 2026, par virement SEPA instantané", "53 million users in 2026, via instant SEPA transfers"),
    (54, "Alternatives souveraines", "Section 2 — note WERO", "Non souverain pour les données stockées sur le cloud états-unien", "Not sovereign if data are stored on US cloud infrastructure."),
    (55, "Alternatives souveraines", "Section 2 — description WERO", "Déjà opérationnel entre particuliers", "Already operational for peer-to-peer payments"),
    (56, "Alternatives souveraines", "Section 2 — réseaux US", "Visa, Mastercard", "Visa, Mastercard"),
    (57, "Alternatives souveraines", "Section 2 — stat réseaux US", "72 % des paiements par carte en zone euro transitent par des schémas états-uniens", "72% of card payments in the eurozone go through US-based schemes"),
    (58, "Alternatives souveraines", "Section 2 — description réseaux US", "Définissent les règles du réseau et autorisent chaque transaction par carte en Europe", "They define network rules and authorise every card transaction in Europe"),
    (59, "Alternatives souveraines", "Section 2 — réseaux EU", "CB, Multibanco, Girocard, Bancontact", "CB, Multibanco, Girocard, Bancontact"),
    (60, "Alternatives souveraines", "Section 2 — stat réseaux EU (coût)", "Jusqu'à 10 x moins chers", "Up to 10x cheaper"),
    (61, "Alternatives souveraines", "Section 2 — stat réseaux EU (fraude)", "Jusqu'à 3 x moins fraudés", "Up to 3x lower fraud rates"),
    (62, "Alternatives souveraines", "Section 2 — description réseaux EU", "Frais de réseaux inférieurs et données traitées et conservées en Europe", "Lower network fees, with data processed and stored in Europe."),
    (63, "Alternatives souveraines", "Section 3 — titre", "LES BANQUES COMMUNIQUENT", "BANKS COMMUNICATE"),
    (64, "Alternatives souveraines", "Section 3 — SWIFT", "SWIFT arsenalisé par section 311 du PATRIOT Act", "SWIFT weaponised under Section 311 of the PATRIOT Act"),
    (65, "Alternatives souveraines", "Section 3 — stat SWIFT", "3e activation de SWIFT comme outil de sanction en dix ans : Iran (2012), Corée du Nord (2017), Russie (2022)", "Third use of SWIFT as a sanctions tool in ten years: Iran (2012), North Korea (2017), Russia (2022)"),
    (66, "Alternatives souveraines", "Section 3 — description SWIFT", "Réseau belge de messagerie financière entre banques exposé à des pressions extraterritoriales américaines", "Belgian interbank financial messaging network exposed to US extraterritorial pressure."),
    (67, "Alternatives souveraines", "Section 4 — titre", "LE RÈGLEMENT EST EFFECTUÉ", "SETTLEMENT IS CARRIED OUT"),
    (68, "Alternatives souveraines", "Section 4 — acteur US", "CHIPS (The Clearing House)", "CHIPS (The Clearing House)"),
    (69, "Alternatives souveraines", "Section 4 — stat CHIPS", "8,9 milliards de dollars Amende infligée à BNP Paribas pour des transactions en dollars réalisées hors du territoire américain", "USD 8.9 billion fine imposed on BNP Paribas for US dollar transactions conducted outside US territory"),
    (70, "Alternatives souveraines", "Section 4 — description CHIPS", "Toute transaction en dollars est soumise à la juridiction états-unienne", "Every US dollar transaction is subject to US jurisdiction"),
    (71, "Alternatives souveraines", "Section 4 — alternative EU", "T2 / TIPS", "T2 / TIPS"),
    (72, "Alternatives souveraines", "Section 4 — stat T2/TIPS", "2 200 milliards réglés par jour en euro, en monnaie de banque centrale, en instantané 24h/24", "EUR 2,200 billion settled daily in euros, in central bank money, instantaneously 24/7"),
    (73, "Alternatives souveraines", "Section 4 — description T2/TIPS", "Un paiement en euro ne dépend jamais de CHIPS", "A euro payment does not depend on CHIPS."),
    (74, "Alternatives souveraines", "Encadré latéral — titre", "QUI EST CONCERNÉ ?", "WHO IS AFFECTED?"),
    (75, "Alternatives souveraines", "Encadré latéral — texte", "Opérateurs d'importance vitale (hôpitaux, énergie, télécommunications, PME, citoyens)", "Operators of vital importance (hospitals, energy, telecommunications, SMEs, citizens)"),
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(table) -> None:
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9E2F3")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)


def build_document(output_path: str) -> None:
    doc = Document()

    title = doc.add_heading("Traductions infographies IRIS — Rapport CB (EN)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Version anglaise relue et homogénéisée — prête pour intégration dans les infographies."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    headers = ["#", "Infographie", "Élément", "Texte original (FR)", "English translation"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_header_row(table)

    col_widths = [Inches(0.35), Inches(1.2), Inches(1.5), Inches(2.5), Inches(2.5)]
    for row_idx, row_data in enumerate(ROWS, start=1):
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                if col_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if row_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F2F2F2")

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Notes : anglais britannique (authorise, weaponised). "
        "Noms propres inchangés : WERO, SecNumCloud, CHIPS, T2, TIPS, CB."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    doc.save(output_path)
    print(f"Document créé : {output_path}")


if __name__ == "__main__":
    build_document(
        r"d:\Users\Proprietaire\Bureau\OLJ\Sahten\Traductions_infographies_IRIS_CB_EN.docx"
    )
